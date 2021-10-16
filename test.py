import docx 
import requests
document = docx.Document("/Users/snoopbob/Downloads/Document.docx")

body = ''
i = 0
while i < len(document.paragraphs):
    body = body + document.paragraphs[i].text
    i = i + 1

response = requests.post('https://api.smmry.com/?SM_API_KEY=7FB201A31A', body.encode('utf-8'))
print(response.json())

# API Key: 7FB201A31A