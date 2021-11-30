import os
import os.path
import pathlib
import time

import docx
import eyed3
import requests


def get_all_file_path(my_path):
    try:
        only_file_names = [f for f in os.listdir(my_path) if os.path.isfile(os.path.join(my_path, f))]
        i = 0
        all_file_path = []
        while i < len(only_file_names):
            file_path = "%s/%s" % (my_path, only_file_names[i])
            all_file_path.append(file_path)
            i = i + 1
        return all_file_path
    except:
        print("Cannot list files from %s" % (my_path))
        return []


def get_my_user_folder_path():
    return pathlib.Path.home()


def move_file(source_filepath: str, destination_folder: str) -> bool:
    try:
        filename = os.path.basename(source_filepath)
        name, extension = os.path.splitext(filename)
        append_name = 1
        destination_filepath = "%s/%s%s" % (destination_folder, name, extension)
        os.makedirs(os.path.dirname(destination_filepath), exist_ok=True)
        while os.path.exists(destination_filepath):
            append_name += 1
            destination_filepath = "%s/%s %d%s" % (destination_folder, name, append_name, extension)
        os.rename(source_filepath, destination_filepath)
        print("Moved %s -> %s" % (source_filepath, destination_filepath))
        return True
    except:
        print("Cannot move %s to %s" % (source_filepath, destination_filepath))
        return False


def change_forbidden_filename(filename: str) -> str:
    forbidden = [
        '\\',
        '/',
        ':',
        '*',
        '?',
        '"',
        '<',
        '>',
        '|',
    ]
    for c in forbidden:
        filename = filename.replace(c, '-')
    return filename


def music(source_file_path: str) -> bool:
    try:
        audio_file = eyed3.load(source_file_path)
        if audio_file.tag == None:
            pass
        else:
            if audio_file.tag.album_artist == None:
                audio_file.tag.album_artist = 'unknown_artist'
            else:
                audio_file.tag.album_artist = change_forbidden_filename(audio_file.tag.album_artist)
            if audio_file.tag.album == None:
                audio_file.tag.album = 'unknown_album'
            else:
                audio_file.tag.album = change_forbidden_filename(audio_file.tag.album)
            artist_filepath = '%s/Music/%s/%s' % (get_my_user_folder_path(), audio_file.tag.album_artist, audio_file.tag.album)
            return move_file(source_file_path, artist_filepath)
    except:
        print("Cannot read %s" % (source_file_path))
        return False


def video(source_file_path: str) -> bool:
    video_mp4_filepath = '%s/Videos' % (get_my_user_folder_path())
    return move_file(source_file_path, video_mp4_filepath)


def others(source_file_path: str) -> bool:
    others_folder = '%s/Others' % (get_my_user_folder_path())
    return move_file(source_file_path, others_folder)


def images(source_file_path: str) -> bool:
    try:
        api_key = 'acc_b580a33df0d99bb'
        api_secret = '508be5bf54a6c634ee73057797de606f'
        upload_info = requests.post(
            'https://api.imagga.com/v2/uploads',
            auth=(api_key, api_secret),
            files={'image': open(source_file_path, 'rb')},
        )
        upload_info = upload_info.json()
        upload_id = upload_info["result"]["upload_id"]
        image_url = 'https://api.imagga.com/v2/categories/personal_photos?image_upload_id=%s' % (upload_id)
        image_catagory = requests.post(
            image_url,
            auth=(api_key, api_secret),
            files={'image': open(source_file_path, 'rb')},
        )
        image_catagory = image_catagory.json()
        image_catagorization = image_catagory["result"]["categories"][0]["name"]["en"]
        destination_image_path = "%s/Pictures/%s" % (get_my_user_folder_path(), image_catagorization)
        return move_file(source_file_path, destination_image_path)
    except:
        print("Cannot access https://api.imagga.com or load %s" % (source_file_path))
        return False


def read_text_file(source_file_path: str) -> str:
    with open(source_file_path) as text:
        return text.readlines()


def read_msword(source_file_path: str) -> str:
    document = docx.Document(source_file_path)
    document_content = ''
    i = 0
    while i < len(document.paragraphs):
        document_content = document_content + document.paragraphs[i].text
        i = i + 1
    return document_content


def categorize_keywords(keywords_list):
    math_science_keywords = [
        'angle',
        'triangle',
        'sqaure',
        'roots',
        'force',
        'mixture',
        'solution',
        'solvent',
        'reaction',
        'equilibrium',
        'speed',
        'acceleration',
        'atom',
        'equation',
        'differenciation',
        'integration',
    ]
    i = 0
    while i < len(keywords_list):
        if keywords_list[i] in math_science_keywords:
            return 'math_science'
        i = i + 1
    return 'everything_else'


smmry_upload_timestamp = 0


def document(source_file_path):
    global smmry_upload_timestamp
    split_filename = os.path.splitext(source_file_path)
    file_ext = str.lower(split_filename[1])
    ms_ext = [
        '.doc',
        '.docx',
        '.ods',
        '.odt',
    ]
    if file_ext == '.txt':
        text_data = read_text_file(source_file_path)
    elif file_ext in ms_ext:
        text_data = read_msword(source_file_path)
    timeout = smmry_upload_timestamp + 10 - time.time()
    if timeout > 0:
        print("Estimated %f s processing document contents %s" % (timeout, source_file_path))
        time.sleep(timeout)
    sm_api_key = '7FB201A31A'
    request_body = dict()
    request_body['sm_api_input'] = text_data
    url = "https://api.smmry.com?SM_API_KEY=%s&SM_KEYWORD_COUNT=%d" % (sm_api_key, 10)
    response = requests.post(url, request_body)
    smmry_upload_timestamp = time.time()
    dict_response = response.json()
    if dict_response.get('sm_api_error') not in [0, 1, 2, 3]:
        keywords = dict_response['sm_api_keyword_array']
        doc_category = categorize_keywords(keywords)
        destination_document_path = "%s/Documents/%s" % (get_my_user_folder_path(), doc_category)
        move_file(source_file_path, destination_document_path)


def main():
    document_ext = [
        '.doc',
        '.docx',
        '.ods',
        '.odt',
        # '.html',
        # '.htm',
        # '.pdf',
        # '.xls',
        # '.xlsx',
        # '.ppt',
        # '.pptx',
        '.txt',
    ]
    video_ext = [
        '.mp4',
        '.mov',
        '.wmv',
        '.flv',
        '.avi',
        '.avchd',
        '.webm',
        '.mkv',
    ]
    music_ext = [
        '.wav',
        '.mp3',
    ]
    img_ext = [
        '.jpeg',
        '.jpg',
        '.png',
        '.apng',
        # '.gif',
        '.jfif',
        '.pjpeg',
        '.pjp',
    ]
    others_ext = [
        '.zip',
        '.7z',
        '.msi',
        '.exe',
        '.pkg',
    ]
    all_file_path = get_all_file_path('%s/Downloads' % (get_my_user_folder_path()))
    i = 0
    num_organized = 0
    num_not_organized = 0
    while i < len(all_file_path):
        try:
            source_file_path = all_file_path[i]
            split_filename = os.path.splitext(source_file_path)
            file_ext = str.lower(split_filename[1])
            if file_ext in music_ext:
                music(source_file_path)
                num_organized += 1
            elif file_ext in video_ext:
                video(source_file_path)
                num_organized += 1
            elif file_ext in img_ext:
                images(source_file_path)
                num_organized += 1
            elif file_ext in others_ext:
                others(source_file_path)
                num_organized += 1
            elif file_ext in document_ext:
                document(source_file_path)
                num_organized += 1
            else:
                num_not_organized += 1
                print("Ignoring %s" % (source_file_path))
        except:
            num_not_organized += 1
            print("Ignoring %s" % (source_file_path))
        i = i + 1
    print("%d files organized" % (num_organized))
    print("%d files left untouched" % (num_not_organized))
    input('Press enter to quit...')


if __name__ == '__main__':
    main()
